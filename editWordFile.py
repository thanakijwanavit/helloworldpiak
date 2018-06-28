#!/usr/bin/env python3
# coding: utf-8

# In[198]:



import docx,requests


# In[199]:


#use data from google sheet to push into word file
def changetext(textToChange,row):
    for i in range(0,len(word.paragraphs)):
        word.paragraphs[i].text=word.paragraphs[i].text.replace(textToChange,s[row][1])
        #word.paragraphs[i].text=word.paragraphs[i].text.replace('----- executive summary----variable----',s[1][1])


# In[200]:


#load word file
word=docx.Document("fan performance report 18_ deckerform.docx")


# In[201]:


#get data from google sheets
r = requests.get('https://docs.google.com/spreadsheets/d/e/2PACX-1vRKEwATIcdsaBDoPYB9P2_td7pKShuZXq767oCbNuyh2P4_h6z02tiVz3469HX8Hm5c1RkjUSDZgP8P/pub?gid=1126628758&single=true&output=csv').text.replace('\r','').split('\n')
s = []
for i in range(0,len(r)):
    s.append(r[i].split(','))
    
print (s)


# In[202]:


for i in range(0,len(s)):
    changetext(s[i][2],i)


# In[203]:


#save file to the specified name
word.save('test2.docx')


# In[204]:


for i in range(0,len(s)):
    print('the response was updated at %s with the value %s' % (s[i][2],s[i][1]))

